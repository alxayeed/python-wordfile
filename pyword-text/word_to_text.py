import docx
import logging

logger = logging.getLogger('py-word   ')
logger.setLevel(level=logging.DEBUG)


def get_paragraphs(file_path):
    doc_file = docx.Document(file_path)
    paragraphs = doc_file.paragraphs
    return paragraphs


def read_paragraphs(paragraphs):
    # get number of paragraph in the word file
    logging.info(f'Total paragraph : {len(paragraphs)}')
    # print all paragraphs text
    for p in paragraphs:
        logger.info(p.text)


def read_runs(paragraphs):
    # Runs means different format of text, eg- normal, bold, italic etc
    logger.info(paragraphs[2].runs[1].text)


def read_all_text(file_path):
    paragraph_list = get_paragraphs(file_path)

    text = []
    for p in paragraph_list:
        text.append(p.text)

    result = "\n".join(text)
    return result


def read_tables(file_path):
    wordDoc = docx.Document(file_path)

    for table in wordDoc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)


if __name__ == '__main__':
    file = 'file/demo.docx'
    # paragraphs = get_paragraphs(file)
    # text = read_all_text(file)

    read_tables(file)
